"""
docx_report_builder.py - Build the MEPI research report as an editable Word document.

Uses python-docx to create a professionally formatted .docx file that mirrors
the structure of the PDF report.  The Word document is suitable for further
editing and can be converted to PDF via Word or LibreOffice.

Usage
-----
    from docx_report_builder import DOCXReportBuilder
    builder = DOCXReportBuilder(
        results_df=results_df,
        chart_paths=chart_paths,
        spatial_map_paths=spatial_map_paths,
        bib_manager=bib_manager,
    )
    path = builder.build("output/report.docx")
"""

from __future__ import annotations

import os
from typing import Dict, List, Optional

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from report_config import (
    REPORT_TITLE,
    REPORT_SUBTITLE,
    REPORT_AUTHORS,
    REPORT_ORGANIZATION,
    REPORT_DATE,
    REPORT_YEAR,
    COLOR_PRIMARY,
    COLOR_SECONDARY,
    COLOR_ACCENT,
    FIGURE_MAX_WIDTH,
    IMAGE_DPI,
)
from report_template import (
    _hex_to_rgb,
    apply_docx_heading_style,
    apply_docx_body_style,
    apply_docx_caption_style,
)

# Max image width in inches for DOCX (6 inches fits A4 with standard margins)
_IMG_MAX_INCHES = 6.0


class DOCXReportBuilder:
    """
    Build the full MEPI DOCX report.

    Parameters
    ----------
    results_df : pd.DataFrame
        Full MEPI results from ``MEPICalculator.calculate()``.
    chart_paths : dict
        Mapping of chart name → PNG file path (from ``ChartGenerator``).
    spatial_map_paths : dict
        Mapping of map name → PNG file path.
    bib_manager : BibliographyManager, optional
        An initialised ``BibliographyManager`` instance.
    """

    def __init__(
        self,
        results_df: pd.DataFrame,
        chart_paths: Optional[Dict[str, str]] = None,
        spatial_map_paths: Optional[Dict[str, str]] = None,
        bib_manager=None,
    ):
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )

        self.df = results_df.copy()
        self.charts = chart_paths or {}
        self.maps = spatial_map_paths or {}
        self.bib = bib_manager
        self._fig_counter = 0
        self._tbl_counter = 0
        self._doc: Optional[Document] = None

    # ------------------------------------------------------------------
    # Counter helpers
    # ------------------------------------------------------------------

    def _next_fig(self) -> str:
        self._fig_counter += 1
        return f"Figure {self._fig_counter}"

    def _next_tbl(self) -> str:
        self._tbl_counter += 1
        return f"Table {self._tbl_counter}"

    def _cite(self, key: str) -> str:
        if self.bib:
            return self.bib.cite(key)
        return f"({key})"

    # ------------------------------------------------------------------
    # Docx helpers
    # ------------------------------------------------------------------

    def _add_heading(self, text: str, level: int = 1):
        heading = self._doc.add_heading(text, level=level)
        apply_docx_heading_style(heading, level=level)
        return heading

    def _add_paragraph(self, text: str):
        para = self._doc.add_paragraph(text)
        apply_docx_body_style(para)
        return para

    def _add_bullet(self, text: str):
        para = self._doc.add_paragraph(text, style="List Bullet")
        for run in para.runs:
            run.font.size = Pt(11)
        return para

    def _add_caption(self, text: str):
        para = self._doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_docx_caption_style(para)
        return para

    def _add_spacer(self):
        self._doc.add_paragraph("")

    def _add_page_break(self):
        self._doc.add_page_break()

    def _add_figure(
        self,
        path: Optional[str],
        caption: str,
        placeholder_label: Optional[str] = None,
        max_width_inches: float = _IMG_MAX_INCHES,
    ):
        """Embed an image or add a placeholder paragraph + caption."""
        fig_label = self._next_fig()
        full_caption = f"{fig_label}: {caption}"

        if path and os.path.isfile(path):
            try:
                # Determine image dimensions and scale
                from PIL import Image as PILImage
                with PILImage.open(path) as pil_img:
                    w_px, h_px = pil_img.size
                dpi = IMAGE_DPI
                w_in = w_px / dpi
                h_in = h_px / dpi
                scale = min(max_width_inches / w_in, 1.0)
                w_final = w_in * scale
                self._doc.add_picture(path, width=Inches(w_final))
                # Centre the picture paragraph
                last_para = self._doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                # Fallback: try with fixed width
                try:
                    self._doc.add_picture(path, width=Inches(min(max_width_inches, 6.0)))
                    self._doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    self._add_placeholder(placeholder_label or f"[Image: {caption}]")
        else:
            self._add_placeholder(placeholder_label or f"[PLACEHOLDER: {caption}]")

        self._add_caption(full_caption)

    def _add_placeholder(self, text: str):
        para = self._doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(150, 150, 150)
            run.font.size = Pt(10)
        # Add a light border box via direct XML (best-effort)
        try:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            for side in ("top", "left", "bottom", "right"):
                bdr = OxmlElement(f"w:{side}")
                bdr.set(qn("w:val"), "single")
                bdr.set(qn("w:sz"), "4")
                bdr.set(qn("w:space"), "4")
                bdr.set(qn("w:color"), "CCCCCC")
                pBdr.append(bdr)
            pPr.append(pBdr)
        except Exception:
            pass
        return para

    def _add_df_table(self, df: pd.DataFrame, caption: Optional[str] = None):
        """Insert a pandas DataFrame as a Word table with header styling."""
        if caption:
            tbl_label = self._next_tbl()
            cap_para = self._doc.add_paragraph(f"{tbl_label}: {caption}")
            cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cap_para.runs:
                run.bold = True
                run.font.size = Pt(10)

        tbl = self._doc.add_table(rows=1, cols=len(df.columns))
        tbl.style = "Table Grid"

        # Header row
        hdr = tbl.rows[0].cells
        primary_rgb = _hex_to_rgb(COLOR_PRIMARY)
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            # Cell shading
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), COLOR_PRIMARY.lstrip("#"))
            tcPr.append(shd)

        # Data rows
        for row_idx, row_data in df.iterrows():
            row = tbl.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(round(val, 4) if isinstance(val, float) else val)
                for run in row[i].paragraphs[0].runs:
                    run.font.size = Pt(9)

        self._doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _cover_page(self):
        doc = self._doc
        doc.add_paragraph("")
        doc.add_paragraph("")
        title_para = doc.add_paragraph(REPORT_TITLE)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.bold = True
            run.font.size = Pt(28)
            run.font.color.rgb = RGBColor(*_hex_to_rgb(COLOR_PRIMARY))

        sub_para = doc.add_paragraph(REPORT_SUBTITLE)
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub_para.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(*_hex_to_rgb(COLOR_SECONDARY))

        doc.add_paragraph("")
        hr_para = doc.add_paragraph("─" * 50)
        hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        for line in [", ".join(REPORT_AUTHORS), REPORT_ORGANIZATION, REPORT_DATE]:
            meta = doc.add_paragraph(line)
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in meta.runs:
                run.font.size = Pt(12)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Table of contents (static – Word will update on open)
    # ------------------------------------------------------------------

    def _table_of_contents(self):
        doc = self._doc
        self._add_heading("Table of Contents", level=1)

        chapters = [
            ("Executive Summary", ""),
            ("Chapter 1: Introduction", ""),
            ("Chapter 2: Methodology", ""),
            ("Chapter 3: Spatial Distribution of Energy Poverty", ""),
            ("Chapter 4: Dimensional Analysis", ""),
            ("Chapter 5: Regional Analysis", ""),
            ("Chapter 6: Temporal Trends", ""),
            ("Chapter 7: Vulnerability Assessment", ""),
            ("Chapter 8: Key Findings", ""),
            ("Chapter 9: Policy Implications and Recommendations", ""),
            ("Chapter 10: Conclusion", ""),
            ("References", ""),
        ]
        for title, _ in chapters:
            para = doc.add_paragraph(title)
            para.style = "List Number"
            for run in para.runs:
                run.font.size = Pt(11)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Executive summary
    # ------------------------------------------------------------------

    def _executive_summary(self):
        self._add_heading("Executive Summary", level=1)
        mean_mepi = round(self.df["mepi_score"].mean(), 3)
        n = len(self.df)
        self._add_paragraph(
            f"This report presents a comprehensive spatial and dimensional analysis of "
            f"energy poverty across Bangladesh using the Multidimensional Energy Poverty "
            f"Index (MEPI) {self._cite('nussbaumer_2011')}. The analysis covers {n} upazilas "
            f"with a mean MEPI score of {mean_mepi:.3f}. Five dimensions—Availability, "
            f"Reliability, Adequacy, Quality, and Affordability—are analysed spatially and "
            f"comparatively. Findings support SDG 7 progress monitoring {self._cite('undp_sdg7_2023')} "
            f"and evidence-based policy targeting."
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 1
    # ------------------------------------------------------------------

    def _chapter_1(self):
        self._add_heading("Chapter 1: Introduction", level=1)

        self._add_heading("1.1 Background", level=2)
        self._add_paragraph(
            f"Energy poverty remains a critical development challenge in Bangladesh "
            f"{self._cite('boardman_1991')}. Despite significant progress in electricity "
            f"access {self._cite('bpdb_2023')}, multidimensional deprivation persists "
            f"across reliability, adequacy, quality, and affordability dimensions "
            f"{self._cite('nussbaumer_2011')}."
        )

        self._add_heading("1.2 Research Objectives", level=2)
        for obj in [
            "Calculate MEPI at the upazila level across Bangladesh.",
            "Identify spatial patterns and hotspots of energy poverty.",
            "Analyse dimensional contributions to overall MEPI scores.",
            "Compare energy poverty across geographic zones and divisions.",
            "Provide evidence-based policy recommendations.",
        ]:
            self._add_bullet(obj)

        self._add_heading("1.3 Literature Review", level=2)
        self._add_paragraph(
            f"The MEPI framework was developed by {self._cite('nussbaumer_2011')} and "
            f"critically reviewed by {self._cite('pelz_2018')}. South Asian applications "
            f"include {self._cite('sadath_2017')}. For Bangladesh, {self._cite('alam_2020')} "
            f"and {self._cite('hossain_2021')} provide country-specific evidence. "
            f"The Alkire–Foster methodology {self._cite('alkire_foster_2011')} underpins "
            f"the counting approach."
        )

        self._add_heading("1.4 Significance", level=2)
        self._add_paragraph(
            f"This study provides upazila-level spatial MEPI estimates aligned with "
            f"Bangladesh's Mujib Climate Prosperity Plan {self._cite('government_bangladesh_2021')} "
            f"and World Bank rural energy programmes {self._cite('world_bank_2022')}."
        )

        self._add_figure(
            self.charts.get("distribution_mepi"),
            "Distribution of MEPI Scores Across All Upazilas",
            "[PLACEHOLDER: Overview MEPI distribution chart]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 2
    # ------------------------------------------------------------------

    def _chapter_2(self):
        self._add_heading("Chapter 2: Methodology", level=1)

        self._add_heading("2.1 MEPI Framework", level=2)
        self._add_paragraph(
            f"The MEPI follows the Alkire–Foster dual-cutoff method "
            f"{self._cite('alkire_foster_2011')}, scoring each upazila on five "
            f"dimensions with equal weights of 0.2 each."
        )

        self._add_heading("2.2 Dimensions and Indicators", level=2)
        dim_df = pd.DataFrame({
            "Dimension": ["Availability", "Reliability", "Adequacy", "Quality", "Affordability"],
            "Indicators": [
                "Electricity access, clean cooking fuel, grid connection",
                "Supply hours, outage frequency, outage duration",
                "Energy consumption, lighting hours, appliance ownership",
                "Voltage fluctuation, satisfaction score, indoor air quality",
                "Expenditure share, cost per kWh, subsidy access",
            ],
            "Weight": ["20%", "20%", "20%", "20%", "20%"],
        })
        self._add_df_table(dim_df, caption="MEPI Dimensions, Indicators, and Weights")

        self._add_heading("2.3 Data Sources", level=2)
        self._add_paragraph(
            f"Data sources: BBS Census {self._cite('bbs_2022')}, BPDB Annual Report "
            f"{self._cite('bpdb_2023')}, SREDA {self._cite('sreda_2022')}, and World Bank "
            f"{self._cite('world_bank_2022')}."
        )

        self._add_figure(
            self.charts.get("dimension_comparison"),
            "Mean Dimension Scores – Methodology Overview",
            "[PLACEHOLDER: Methodology diagram]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 3
    # ------------------------------------------------------------------

    def _chapter_3(self):
        self._add_heading("Chapter 3: Spatial Distribution of Energy Poverty", level=1)

        self._add_paragraph(
            f"Spatial analysis reveals pronounced geographic disparities "
            f"{self._cite('tobler_1970')}. The following maps display upazila-level MEPI "
            f"and dimension scores."
        )

        maps_info = [
            ("mepi_spatial_map", "Overall MEPI Spatial Distribution",
             "[INSERT: mepi_spatial_map.png – Overall MEPI map from ~/spatial_maps_png/]"),
            ("availability_map", "Availability Dimension Map",
             "[INSERT: availability_map.png – from ~/spatial_maps_png/]"),
            ("reliability_map", "Reliability Dimension Map",
             "[INSERT: reliability_map.png – from ~/spatial_maps_png/]"),
            ("adequacy_map", "Adequacy Dimension Map",
             "[INSERT: adequacy_map.png – from ~/spatial_maps_png/]"),
            ("quality_map", "Quality Dimension Map",
             "[INSERT: quality_map.png – from ~/spatial_maps_png/]"),
            ("affordability_map", "Affordability Dimension Map",
             "[INSERT: affordability_map.png – from ~/spatial_maps_png/]"),
        ]

        for key, caption, placeholder in maps_info:
            self._add_figure(self.maps.get(key), caption, placeholder)

        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 4
    # ------------------------------------------------------------------

    def _chapter_4(self):
        self._add_heading("Chapter 4: Dimensional Analysis", level=1)

        dim_cols = [c for c in ["Availability_score", "Reliability_score", "Adequacy_score",
                                  "Quality_score", "Affordability_score"] if c in self.df.columns]
        if dim_cols:
            stats_df = pd.DataFrame({
                "Dimension": [c.replace("_score", "") for c in dim_cols],
                "Mean": self.df[dim_cols].mean().round(4).values,
                "Std": self.df[dim_cols].std().round(4).values,
                "Min": self.df[dim_cols].min().round(4).values,
                "Max": self.df[dim_cols].max().round(4).values,
            })
            self._add_df_table(stats_df, caption="Descriptive Statistics of Dimension Scores")

        self._add_figure(
            self.charts.get("dimension_comparison"),
            "Mean MEPI Dimension Scores Comparison",
            "[PLACEHOLDER: Dimension comparison chart]",
        )

        self._add_heading("4.1 Correlation Analysis", level=2)
        self._add_paragraph(
            f"Pairwise correlations between dimension scores reveal energy poverty "
            f"interdependencies {self._cite('zhang_2019')}."
        )
        self._add_figure(
            self.charts.get("correlation_heatmap"),
            "Dimension Score Correlation Heatmap",
            "[INSERT: Correlation heatmap]",
        )

        self._add_heading("4.2 Dimension Heatmap", level=2)
        self._add_figure(
            self.charts.get("dimension_heatmap"),
            "Dimension Score Heatmap by Upazila",
            "[INSERT: Dimension heatmap]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 5
    # ------------------------------------------------------------------

    def _chapter_5(self):
        self._add_heading("Chapter 5: Regional Analysis", level=1)

        zone_col = None
        for col in ("geographic_zone", "division", "district"):
            if col in self.df.columns:
                zone_col = col
                break

        if zone_col:
            score_cols = ["mepi_score"] + [c for c in self.df.columns if c.endswith("_score") and c != "mepi_score"]
            region_tbl = self.df.groupby(zone_col)[score_cols].mean().round(3).reset_index()
            region_tbl.columns = [c.replace("_score", "").replace("_", " ").title() for c in region_tbl.columns]
            self._add_df_table(region_tbl, caption=f"Average MEPI Scores by {zone_col.replace('_', ' ').title()}")

        for title, text in [
            ("5.1 Coastal Zone", "Coastal upazilas face energy poverty driven by tidal flooding, saline intrusion, and fragmented grid infrastructure."),
            ("5.2 Char Islands", "River island communities suffer extreme energy isolation with limited grid access and high solar home system dependence."),
            ("5.3 Haor Wetlands", "Seasonal inundation disrupts power supply for up to six months annually."),
            ("5.4 Chittagong Hill Tracts", "Lowest electricity access rates due to rugged terrain and high connection costs."),
            ("5.5 Sundarbans Fringe", "Dual challenges of flood risk and energy isolation compound wood fuel dependence."),
        ]:
            self._add_heading(title, level=2)
            self._add_paragraph(text)

        self._add_figure(
            self.charts.get("regional_comparison"),
            "Regional Comparison of MEPI Dimension Scores",
            "[PLACEHOLDER: Regional comparison chart]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 6
    # ------------------------------------------------------------------

    def _chapter_6(self):
        self._add_heading("Chapter 6: Temporal Trends", level=1)

        self._add_paragraph(
            f"Bangladesh has made significant progress in electricity access since 2010 "
            f"{self._cite('bpdb_2023')}, but improvements in reliability and affordability "
            f"have lagged. Temporal tracking is essential for SDG 7 monitoring "
            f"{self._cite('undp_sdg7_2023')}."
        )

        self._add_figure(
            self.charts.get("temporal_trend"),
            "Illustrative Temporal Trend of Mean MEPI Score (2018–2023)",
            "[PLACEHOLDER: Temporal trend chart – requires multi-year data]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 7
    # ------------------------------------------------------------------

    def _chapter_7(self):
        self._add_heading("Chapter 7: Vulnerability Assessment", level=1)

        self._add_heading("7.1 Hotspot Identification", level=2)
        threshold = self.df["mepi_score"].quantile(0.75)
        hotspot_df = self.df[self.df["mepi_score"] >= threshold].copy()
        show_cols = [c for c in ["upazila_name", "district", "mepi_score", "poverty_category"] if c in hotspot_df.columns]
        if show_cols:
            hs_tbl = hotspot_df[show_cols].sort_values("mepi_score", ascending=False).reset_index(drop=True)
            hs_tbl.columns = [c.replace("_", " ").title() for c in hs_tbl.columns]
            self._add_df_table(hs_tbl, caption="Energy Poverty Hotspot Upazilas (MEPI ≥ 75th Percentile)")

        self._add_figure(
            self.maps.get("hotspot_map"),
            "Energy Poverty Hotspot and Vulnerability Map",
            "[INSERT: hotspot_map.png – Vulnerability clusters]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 8
    # ------------------------------------------------------------------

    def _chapter_8(self):
        self._add_heading("Chapter 8: Key Findings", level=1)

        # Summary statistics
        stats_df = pd.DataFrame({
            "Statistic": ["N (Upazilas)", "Mean MEPI", "Median MEPI", "Std. Dev.", "Min", "Max"],
            "Value": [
                str(len(self.df)),
                f"{self.df['mepi_score'].mean():.4f}",
                f"{self.df['mepi_score'].median():.4f}",
                f"{self.df['mepi_score'].std():.4f}",
                f"{self.df['mepi_score'].min():.4f}",
                f"{self.df['mepi_score'].max():.4f}",
            ],
        })
        self._add_df_table(stats_df, caption="Summary Statistics of MEPI Scores")

        self._add_heading("8.1 Top 10 Most Energy-Poor Upazilas", level=2)
        top10 = self.df.nlargest(10, "mepi_score")
        show_cols = [c for c in ["upazila_name", "district", "mepi_score", "poverty_category"] if c in top10.columns]
        top10_tbl = top10[show_cols].reset_index(drop=True)
        top10_tbl.columns = [c.replace("_", " ").title() for c in top10_tbl.columns]
        self._add_df_table(top10_tbl, caption="Top 10 Most Energy-Poor Upazilas")

        self._add_figure(
            self.charts.get("top10_most_poor"),
            "Top 10 Most Energy-Poor Upazilas – MEPI Bar Chart",
            "[PLACEHOLDER: Top 10 most poor bar chart]",
        )

        self._add_heading("8.2 Top 10 Least Energy-Poor Upazilas", level=2)
        bot10 = self.df.nsmallest(10, "mepi_score")
        bot10_tbl = bot10[show_cols].reset_index(drop=True)
        bot10_tbl.columns = [c.replace("_", " ").title() for c in bot10_tbl.columns]
        self._add_df_table(bot10_tbl, caption="Top 10 Least Energy-Poor Upazilas")

        self._add_figure(
            self.charts.get("top10_least_poor"),
            "Top 10 Least Energy-Poor Upazilas – MEPI Bar Chart",
            "[PLACEHOLDER: Top 10 least poor bar chart]",
        )

        self._add_heading("8.3 Category Distribution", level=2)
        self._add_figure(
            self.charts.get("poverty_category_pie"),
            "Poverty Category Distribution",
            "[PLACEHOLDER: Poverty category pie chart]",
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 9
    # ------------------------------------------------------------------

    def _chapter_9(self):
        self._add_heading("Chapter 9: Policy Implications and Recommendations", level=1)

        recs = [
            ("9.1 Grid Extension and Mini-Grids",
             f"Prioritise grid extension to the top 50 most energy-poor upazilas. "
             f"Deploy solar mini-grids where grid extension is not viable {self._cite('irena_2023')}."),
            ("9.2 Clean Cooking Fuel",
             f"Scale up LPG subsidies and improved cook stoves in coastal and haor upazilas {self._cite('sreda_2022')}."),
            ("9.3 Reliability Improvement",
             "Invest in distribution network upgrades and smart metering in high-outage areas."),
            ("9.4 Affordability Interventions",
             f"Expand lifeline tariffs and subsidy programmes {self._cite('world_bank_2022')}."),
            ("9.5 Data and Monitoring",
             f"Establish an annual energy poverty monitoring system aligned with MEPI dimensions {self._cite('bbs_2022')}."),
        ]
        for title, text in recs:
            self._add_heading(title, level=2)
            self._add_paragraph(text)

        self._add_page_break()

    # ------------------------------------------------------------------
    # Chapter 10
    # ------------------------------------------------------------------

    def _chapter_10(self):
        self._add_heading("Chapter 10: Conclusion", level=1)

        self._add_paragraph(
            f"This study has provided a comprehensive MEPI analysis of energy poverty "
            f"in Bangladesh {self._cite('nussbaumer_2011')}, revealing marked geographic "
            f"disparities and multidimensional deprivation patterns. The upazila-level "
            f"spatial estimates enable targeted policy interventions aligned with SDG 7 "
            f"{self._cite('undp_sdg7_2023')}."
        )

        self._add_heading("10.1 Future Research", level=2)
        for direction in [
            "Longitudinal MEPI tracking using multi-wave panel data.",
            "Climate vulnerability integration into the MEPI framework.",
            "Micro-level analysis using satellite-derived night-light data.",
            "Gender-disaggregated energy poverty analysis.",
        ]:
            self._add_bullet(direction)

        self._add_heading("10.2 Limitations", level=2)
        self._add_paragraph(
            "The analysis is limited to 20 sample upazilas; a nationally representative "
            "sample would provide more definitive spatial estimates. Equal weighting of "
            "dimensions may not reflect local priorities."
        )
        self._add_page_break()

    # ------------------------------------------------------------------
    # References
    # ------------------------------------------------------------------

    def _references(self):
        self._add_heading("References", level=1)

        if self.bib:
            bib_list = self.bib.bibliography_list(only_cited=True)
            for entry in bib_list:
                para = self._doc.add_paragraph(entry)
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.first_line_indent = Inches(-0.25)
                for run in para.runs:
                    run.font.size = Pt(10)
                self._doc.add_paragraph("")
        else:
            self._add_paragraph("[Bibliography not available]")

    # ------------------------------------------------------------------
    # Main build method
    # ------------------------------------------------------------------

    def build(self, output_path: str) -> str:
        """
        Compose and write the full DOCX report.

        Parameters
        ----------
        output_path : str
            Destination file path for the DOCX.

        Returns
        -------
        str
            Absolute path of the saved DOCX file.
        """
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        self._doc = Document()

        # Page setup – A4
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE

        section = self._doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

        # Set default body font
        style = self._doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Build content
        self._cover_page()
        self._table_of_contents()
        self._executive_summary()
        self._chapter_1()
        self._chapter_2()
        self._chapter_3()
        self._chapter_4()
        self._chapter_5()
        self._chapter_6()
        self._chapter_7()
        self._chapter_8()
        self._chapter_9()
        self._chapter_10()
        self._references()

        self._doc.save(output_path)
        abs_path = os.path.abspath(output_path)
        print(f"DOCX report saved: {abs_path}")
        return abs_path
